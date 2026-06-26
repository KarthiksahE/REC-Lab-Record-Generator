from docxtpl import DocxTemplate
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import uuid
import copy
import logging
from typing import Dict, Any, List
from app.config import settings
from app.utils.qr_generator import generate_qr_code
from app.utils.pdf_converter import convert_docx_to_pdf

logger = logging.getLogger(__name__)

class DocumentService:
    def __init__(self):
        # Resolve absolute template path
        base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        self.template_path = os.path.join(base_dir, settings.TEMPLATE_PATH)
        self.temp_dir = os.path.join(base_dir, "generated", "temp")
        os.makedirs(self.temp_dir, exist_ok=True)

    def generate_document(self, data: Dict[str, Any], save_persistently: bool = False) -> Dict[str, str]:
        """
        Loads the lab record template, replaces fields, expands experiments index,
        embeds QR codes, and compiles PDF.
        
        Args:
            data (dict): Combined form data containing document info and experiments list.
            save_persistently (bool): Whether files should be stored long-term (or deleted after response).
            
        Returns:
            dict: Paths to generated DOCX and PDF files.
        """
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"Base Word template not found at: {self.template_path}")
            
        unique_id = str(uuid.uuid4())
        docx_out_path = os.path.join(self.temp_dir, f"record_{unique_id}.docx")
        
        # Lists to keep track of temp QR image paths for deletion later
        temp_qr_paths = []
        
        try:
            # 1. Load the template
            doc = DocxTemplate(self.template_path)
            doc.init_docx()
            
            # Autocorrect template placeholder typos (e.g. {{SubjectCode})
            self._autocorrect_template_placeholders(doc)
            
            # 2. Extract experiments data
            experiments_list = data.get("experiments", [])
            
            # 3. Handle index table row duplication
            self._process_experiment_table(doc, experiments_list, temp_qr_paths)
            
            # 4. Prepare context for general page-level placeholders
            # Map frontend names to XML placeholder names
            context = {
                "StudentName": data.get("student_name", ""),
                "RegisterNumber": data.get("register_number", ""),
                "SubjectCode": data.get("course_code", ""),
                "SubjectName": data.get("course_name", ""),
                "AcademicYear": data.get("academic_year", ""),
                "semester": data.get("semester", ""),
                "year": data.get("year", ""),
                "branch": data.get("department", ""), # branch maps to Department in template
                
                # Additional fields (can be utilized if template includes them)
                "Department": data.get("department", ""),
                "Year": data.get("year", ""),
                "Semester": data.get("semester", ""),
                "Faculty": data.get("faculty", ""),
                "LabName": data.get("lab_name", ""),
                "Institution": data.get("institution", "")
            }
            
            # 5. Render general placeholders using docxtpl
            doc.render(context)
            
            # Fix page borders compatibility for LibreOffice conversion
            self._add_page_borders(doc)
            
            # Save the generated docx file
            doc.save(docx_out_path)
            logger.info(f"Generated DOCX saved to: {docx_out_path}")
            
            # 6. Convert DOCX to PDF using LibreOffice
            pdf_out_path = convert_docx_to_pdf(docx_out_path, self.temp_dir)
            logger.info(f"Generated PDF saved to: {pdf_out_path}")
            
            return {
                "docx_path": docx_out_path,
                "pdf_path": pdf_out_path,
                "unique_id": unique_id,
                "temp_qrs": temp_qr_paths
            }
            
        except Exception as e:
            # Clean up any generated assets immediately on error
            if os.path.exists(docx_out_path):
                try:
                    os.remove(docx_out_path)
                except Exception:
                    pass
            for qr_path in temp_qr_paths:
                if os.path.exists(qr_path):
                    try:
                        os.remove(qr_path)
                    except Exception:
                        pass
            logger.error(f"Error during document generation: {str(e)}")
            raise e

    def _process_experiment_table(self, doc: DocxTemplate, experiments: List[Dict[str, Any]], temp_qr_paths: List[str]):
        """
        Locates the index table, duplicates the placeholder row for each experiment,
        inserts serial number, date, title, and generates/embeds the transparent QR code image.
        """
        # Find the index table
        target_table = None
        for table in doc.tables:
            # Check if this table has our headers (e.g. S.No, Date, Name of the Experiment, etc.)
            if len(table.rows) > 0:
                first_row_text = "".join(cell.text for cell in table.rows[0].cells).lower()
                if "s.no" in first_row_text and "experiment" in first_row_text:
                    target_table = table
                    break
                    
        if not target_table:
            logger.warning("Index table was not found in the template. Skipping table row duplication.")
            return

        # Locate the template rows (rows containing placeholders) and remove them
        # Let's inspect rows starting from row 1 (row 0 is header)
        template_row = None
        rows_to_remove = []
        
        for idx in range(1, len(target_table.rows)):
            row = target_table.rows[idx]
            row_text = "".join(cell.text for cell in row.cells)
            
            # Check if this row is a placeholder row
            if "[Experiment Title Placeholder]" in row_text or "[Exp No]" in row_text:
                if template_row is None:
                    # Keep the first placeholder row as our clone source
                    template_row = row
                rows_to_remove.append(row)

        if not template_row:
            logger.warning("No experiment placeholder rows found in the index table. Row duplication skipped.")
            return

        # If there are no experiments, we'll just clear/remove the placeholder rows
        # If there are experiments, we duplicate the template row for each experiment
        if experiments:
            for s_no_idx, exp in enumerate(experiments, start=1):
                # 1. Duplicate the row XML properties and structure
                new_tr = copy.deepcopy(template_row._tr)
                target_table._tbl.append(new_tr)
                
                # Wrap the XML element in python-docx _Row class
                from docx.table import _Row
                new_row = _Row(new_tr, target_table)
                
                # Helper to write to cell while preserving template styling and alignments
                def format_and_write_cell(cell, text):
                    cell.vertical_alignment = 1 # Center vertically
                    
                    # Clear extra paragraphs
                    while len(cell.paragraphs) > 1:
                        p_element = cell.paragraphs[-1]._element
                        p_element.getparent().remove(p_element)
                        
                    p = cell.paragraphs[0]
                    if p.runs:
                        p.runs[0].text = text
                        # Clear text from any other runs inside this paragraph
                        for r in p.runs[1:]:
                            r.text = ""
                    else:
                        p.text = text
                
                # 2. Fill the cell values
                # Cell 0: S.No (or Exp No)
                format_and_write_cell(new_row.cells[0], str(s_no_idx))
                
                # Cell 1: Date
                date_val = exp.get("date", "")
                format_and_write_cell(new_row.cells[1], date_val if date_val else "")
                
                # Cell 2: Experiment Title
                format_and_write_cell(new_row.cells[2], exp.get("title", ""))
                
                # Cell 3: QR Code (from GitHub link)
                new_row.cells[3].vertical_alignment = 1 # Center vertically
                github_url = exp.get("github_url", "")
                if github_url:
                    qr_filename = f"qr_{uuid.uuid4()}.png"
                    qr_path = os.path.join(self.temp_dir, qr_filename)
                    
                    try:
                        # Generate transparent PNG QR code
                        generate_qr_code(github_url, qr_path)
                        temp_qr_paths.append(qr_path)
                        
                        # Clear cell text/runs and insert the image
                        p = new_row.cells[3].paragraphs[0]
                        for r in p.runs:
                            r.text = ""
                        run = p.add_run()
                        # Add image and scale to fit inside cell without stretching
                        run.add_picture(qr_path, width=Inches(0.85))
                    except Exception as qr_err:
                        logger.error(f"Failed to insert QR code for row {s_no_idx}: {str(qr_err)}")
                        format_and_write_cell(new_row.cells[3], "QR Error")
                else:
                    format_and_write_cell(new_row.cells[3], "")
                    
                # Cell 4: Signature (Leave blank space)
                format_and_write_cell(new_row.cells[4], "")

        # Remove the original template rows
        for row in rows_to_remove:
            target_table._tbl.remove(row._tr)

    def _add_page_borders(self, doc):
        """
        Fixes page borders compatibility with LibreOffice PDF conversion by ensuring
        existing borders are offset from 'page'.
        """
        for section in doc.sections:
            sectPr = section._sectPr
            pgBorders = sectPr.find(qn('w:pgBorders'))
            if pgBorders is not None:
                # Setting offsetFrom to 'page' ensures LibreOffice renders the border in PDFs
                pgBorders.set(qn('w:offsetFrom'), 'page')
                
            # Ensure the individual page borders have spacing so they don't overlap text
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = pgBorders.find(qn(f'w:{border_name}'))
                if border is not None:
                    # Set default spacing margin (24pt) if not present
                    if not border.get(qn('w:space')):
                        border.set(qn('w:space'), '24')

    def clean_up_files(self, paths_dict: Dict[str, Any]):
        """
        Deletes temporary files (docx, pdf, qrs) generated during a document session.
        """
        # Delete generated DOCX
        docx_path = paths_dict.get("docx_path")
        if docx_path and os.path.exists(docx_path):
            try:
                os.remove(docx_path)
            except Exception as e:
                logger.error(f"Failed to delete temp docx: {str(e)}")
                
        # Delete generated PDF
        pdf_path = paths_dict.get("pdf_path")
        if pdf_path and os.path.exists(pdf_path):
            try:
                os.remove(pdf_path)
            except Exception as e:
                logger.error(f"Failed to delete temp pdf: {str(e)}")
                
        # Delete all QR code images
        temp_qrs = paths_dict.get("temp_qrs", [])
        for qr_path in temp_qrs:
            if os.path.exists(qr_path):
                try:
                    os.remove(qr_path)
                except Exception as e:
                    logger.error(f"Failed to delete temp QR image: {str(e)}")

    def _autocorrect_template_placeholders(self, doc: DocxTemplate):
        """
        Scans all body text, tables, headers, and footers in the template
        and fixes any syntax typos where double braces are single-closed,
        e.g., changing '{{SubjectCode}' to '{{SubjectCode}}'.
        """
        # Scan body paragraphs
        for p in doc.paragraphs:
            self._fix_runs(p.runs)
            
        # Scan tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._fix_runs(p.runs)
                        
        # Scan headers and footers
        for section in doc.sections:
            # Check headers
            headers = [section.header, getattr(section, "first_page_header", None), getattr(section, "even_page_header", None)]
            for header in headers:
                if header:
                    for p in header.paragraphs:
                        self._fix_runs(p.runs)
                    for table in header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    self._fix_runs(p.runs)
                                    
            # Check footers
            footers = [section.footer, getattr(section, "first_page_footer", None), getattr(section, "even_page_footer", None)]
            for footer in footers:
                if footer:
                    for p in footer.paragraphs:
                        self._fix_runs(p.runs)
                    for table in footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    self._fix_runs(p.runs)

    def _fix_runs(self, runs):
        """
        Inspects runs list in a paragraph. If it detects a broken run sequence
        like ['{{', 'SubjectCode', '}'], it replaces the '}' run with '}}'
        to fix Jinja parser errors. It also heals single runs like '{{SubjectCode}'.
        """
        if not runs:
            return
            
        # 1. Fix split run sequence typo: {{ -> Text -> }
        for i in range(len(runs) - 2):
            try:
                # Check for double open brace, text content, and single close brace
                if runs[i].text and "{?" in runs[i].text or "{{" in runs[i].text:
                    if runs[i+2].text == "}":
                        runs[i+2].text = "}}"
            except Exception:
                pass
                
        # 2. Fix single run typo: {{SubjectCode}
        for r in runs:
            try:
                if r.text and "{{" in r.text and "}" in r.text and not "}}" in r.text:
                    clean_text = r.text.strip()
                    if clean_text.endswith("}") and not clean_text.endswith("}}"):
                        r.text = r.text + "}"
            except Exception:
                pass

# Singleton instance
document_service = DocumentService()
